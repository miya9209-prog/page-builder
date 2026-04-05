import base64
import io
import json
import mimetypes
import re
import time
from typing import Any, Dict, List

import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openai import OpenAI, RateLimitError

st.set_page_config(page_title="PAGE BUILDER", layout="wide")

# ─────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────
for key, default in [
    ("reset_nonce", 0),
    ("naming_result", ""),
    ("naming_input_value", ""),
    ("generated_result", ""),
    ("generated_docx", b""),
    ("generated_file_stem", "page_builder"),
]:
    if key not in st.session_state:
        st.session_state[key] = default

st.markdown(
    "<style>div[data-testid='stButton'] > button { min-height: 42px; }</style>",
    unsafe_allow_html=True,
)

st.title("MISHARP PAGE BUILDER")
st.caption("구매전환율 상승을 위한 상세페이지 기획 + 상품 원고 생성기")

api_key = st.secrets.get("OPENAI_API_KEY", "")
if not api_key:
    st.warning("OPENAI_API_KEY가 설정되지 않았습니다. Streamlit Cloud Secrets 또는 .streamlit/secrets.toml을 확인해 주세요.")
    st.stop()

client = OpenAI(api_key=api_key)

# ─────────────────────────────────────────────
# 고정 HTML 헤더
# ─────────────────────────────────────────────
FIXED_HTML_HEAD = (
    '<meta http-equiv="X-UA-Compatible" content="IE=edge,chrome=1">\n'
    '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
    '<link href="http://fonts.googleapis.com/css?family=Roboto" rel="stylesheet" type="text/css">\n'
    '<link href="http://netdna.bootstrapcdn.com/font-awesome/4.3.0/css/font-awesome.min.css" rel="stylesheet" type="text/css">\n'
    '<link href="/SRC2/cssmtmenu/style.css" rel="stylesheet" type="text/css">\n'
    '<link href="//spoqa.github.io/spoqa-han-sans/css/SpoqaHanSans-kr.css" rel="stylesheet" type="text/css">\n'
    '<link href="//misharp.co.kr/subtap.css" rel="stylesheet" type="text/css">'
)

NAME_PROMPT = (
    "너는 4050 여성 패션 쇼핑몰 미샵의 상품 네이밍 전문가다.\n"
    "- 상품 주요 특징을 반영해 상품명을 20개 제안한다.\n"
    "- 각 상품명은 공백 포함 최대 18자 이내.\n"
    "- 반드시 단어와 단어 사이를 자연스럽게 띄어쓴다.\n"
    "- AI 검색, 키워드 검색 모두 고려한다.\n"
    "- 디테일/형태/원단/핏 등을 반영한 단어 + 카테고리명을 포함한다.\n"
    "- 필요하면 세련되고 여성스러운 단어를 앞에 붙여도 된다.\n"
    "- 번호, 설명, 코드펜스 없이 한 줄에 하나씩 20개만 출력한다."
)


# ─────────────────────────────────────────────
# 유틸 함수
# ─────────────────────────────────────────────
def chat_with_retry(*, model, messages, temperature=0.3, max_retries=2):
    last_exc = None
    for attempt in range(max_retries + 1):
        try:
            return client.chat.completions.create(
                model=model, messages=messages, temperature=temperature
            )
        except RateLimitError as exc:
            last_exc = exc
            if attempt >= max_retries:
                raise
            time.sleep(2 * (attempt + 1))
        except Exception as exc:
            last_exc = exc
            if attempt >= max_retries:
                raise
            time.sleep(1 * (attempt + 1))
    raise last_exc


def file_to_content_item(uploaded_file):
    mime = uploaded_file.type or mimetypes.guess_type(uploaded_file.name)[0] or "image/jpeg"
    b64 = base64.b64encode(uploaded_file.read()).decode("utf-8")
    return {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}


def extract_lines_with_digits(text):
    out = []
    for raw in (text or "").splitlines():
        line = raw.strip()
        if line and re.search(r"\d", line):
            out.append(line)
    return out


def combine_measurements(top, bottom, dress):
    lines = []
    for block in [top, bottom, dress]:
        lines.extend(extract_lines_with_digits(block))
    return lines


def count_colors(color_text):
    if not color_text.strip():
        return 0
    # 슬래시·쉼표로 먼저 분리
    text = color_text.replace(" / ", "\n").replace("/", "\n").replace(",", "\n")
    result = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        # "1 블랙 2 네이비 3 베이지" 처럼 숫자가 구분자인 경우 각각 분리
        sub = re.split(r"(?=\s*\d+\s+[가-힣a-zA-Z])", line)
        for s in sub:
            s = re.sub(r"^\s*\d+\s*", "", s).strip()
            if s:
                result.append(s)
    return len(result) if result else 0


def apply_color_count_to_name(product_name, color_text):
    count = count_colors(color_text)
    suffix = f"({count} color)" if count > 0 else "(color)"
    name = (product_name or "").strip()
    name = re.sub(r"\(\s*\d+\s*color\s*\)", suffix, name, flags=re.I)
    name = re.sub(r"\(\s*color\s*\)", suffix, name, flags=re.I)
    return name


def clean_line(line):
    line = (line or "").strip()
    line = re.sub(r"^[\-\u2022\u25aa\u29bf\s]+", "", line)
    line = re.sub(r"\s+", " ", line).strip()
    # ",." 중복 패턴 정리
    line = re.sub(r",\.+", ".", line)
    return line


def ensure_period(line):
    """마침표로 끝나게 보장. 쉼표/쉼표+마침표로 끝나는 경우도 정리."""
    line = clean_line(line)
    if not line:
        return ""
    # 쉼표+마침표 중복 제거: ",." → "."
    line = re.sub(r",\.+$", ".", line)
    # 쉼표로만 끝나는 경우: 쉼표 제거 후 마침표
    if line.endswith(","):
        line = line[:-1].rstrip()
    if re.search(r"[.!?\u2026]$", line):
        return line
    return line + "."


def format_measurement_lines(lines):
    if not lines:
        return "실측사이즈 정보를 입력해 주세요."
    formatted = []
    for line in lines:
        line = re.sub(r"\s+단위:cm$", "", line).strip()
        for label in ["XL", "L", "M", "S", "F"]:
            line = re.sub(rf"\s+{label}\s+", f"<br>{label} ", line)
        line = re.sub(r"\s+", " ", line)
        formatted.append(line)
    return "<br>".join(formatted) + " (단위: cm)"


def format_material_line(material):
    mat = (material or "").strip()
    if not mat:
        return "(건조기사용금지)"
    if "(건조기사용금지)" not in mat:
        mat = mat + " (건조기사용금지)"
    return mat


def result_to_docx_bytes(result_text):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Dotum"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "돋움")
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.5
    for line in result_text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(line)
        run.font.name = "Dotum"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "돋움")
        run.font.size = Pt(10)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def reset_all():
    st.session_state.reset_nonce += 1
    st.session_state.naming_result = ""
    st.session_state.naming_input_value = ""
    st.session_state.generated_result = ""
    st.session_state.generated_docx = b""
    st.session_state.generated_file_stem = "page_builder"


# ─────────────────────────────────────────────
# JSON 파싱
# ─────────────────────────────────────────────
def extract_json(text):
    text = (text or "").strip()
    text = re.sub(r"^```(?:json)?", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    if text.startswith("{") and text.endswith("}"):
        return json.loads(text)
    m = re.search(r"\{[\s\S]*\}", text)
    if not m:
        raise ValueError("JSON 응답을 찾지 못했습니다.")
    return json.loads(m.group(0))


# ─────────────────────────────────────────────
# 프롬프트 빌더
# ─────────────────────────────────────────────
def build_generation_prompt(data, additional_request):
    measurement_str = " / ".join(data["measurement_lines"]) if data["measurement_lines"] else "입력 없음"

    prompt = (
        "너는 10년 이상 4050 여성 패션몰 미샵 상세페이지를 써 온 최고 수준의 한국어 커머스 카피라이터다.\n"
        "실제 고객이 망설이는 이유를 먼저 해소하고 구매를 돕는 고급 원고를 작성한다.\n"
        "SEO/AEO/GEO를 고려해 상품명, 카테고리, 소재, 핏, 활용 장면, 고객 니즈를 자연스럽게 녹인다.\n\n"
        "[중요 원칙]\n"
        "1. 입력 문구를 그대로 복붙하지 말고 자연스럽고 매력적인 문장으로 재작성한다.\n"
        "2. 없는 디테일은 절대 만들지 않는다. 타이/스카프/포켓/단추/절개 등은 입력 근거가 있을 때만 언급한다.\n"
        "3. 문장형은 문장형으로, 리스팅형은 리스팅형으로 작성한다.\n"
        "4. 결과는 반드시 JSON만 출력한다. 코드펜스 절대 금지.\n\n"
        "[추천 블록 작성 규칙 - 절대 준수]\n"
        "recommend_lines 4줄은 반드시 아래 형태로만 작성한다:\n"
        "  예시) 오피스룩부터 하객룩까지 다양하게 활용할 수 있는 블라우스를 찾으시는 분\n"
        "  예시) 군살을 자연스럽게 커버하면서도 여성스러운 실루엣을 원하시는 분\n"
        "  예시) 구김 걱정 없이 하루 종일 깔끔하게 입고 싶으신 분\n"
        "  예시) 브랜드 퀄리티의 고급 소재를 편하게 입고 싶으신 분\n"
        "절대 금지 표현: 추천합니다, 권해드립니다, 적합합니다, 잘 어울립니다, 분께, 고객님께\n"
        "모든 문장은 반드시 '분' 또는 '분.'으로 끝나야 한다.\n\n"
        "[MD원고 작성 규칙]\n"
        "md_sections.purchase_note 와 md_sections.ending 은 반드시 빈 배열 []로 출력한다.\n"
        "[구매 전 꼭 확인해 주세요] 섹션은 절대 작성하지 않는다.\n"
        "마무리 감성 문장은 작성하지 않는다.\n"
        "MD원고는 choice / fabric / fit / occasion 4개 섹션만 작성한다.\n\n"
        "[MD원고 문장 스타일 - 매우 중요]\n"
        "각 섹션의 문장은 짧은 단문을 나열하지 말고, 자연스럽게 이어지는 흐름으로 작성한다.\n"
        "한 줄은 15~25자 내외로, 의미가 연결되는 곳에서 자연스럽게 줄바꿈한다.\n"
        "문장은 자연스러운 구어체로 완성한다. 명사형으로 끝낼 필요는 없으며, 서술형·접속형 모두 허용한다.\n"
        "아래 <style_guide> 안의 예시는 스타일 참고용이다. 이 문장들을 그대로 출력에 포함하지 않는다.\n\n"
        "<style_guide>\n"
        "[choice 예시]\n"
        "미샵의 브랜드 납품의 고급 공정 파트너사 상품,\n"
        "그만큼 자신있는 고퀄리티 상품입니다.\n"
        "원단에서 입어서의 편안함까지\n"
        "브랜드급의 매력적인 블라우스입니다.\n\n"
        "[fabric 예시]\n"
        "겉감과 안감 모두 폴리에스터 100% 소재로\n"
        "가볍고 내구성이 뛰어납니다.\n"
        "구김이 적고 형태가 잘 유지되어\n"
        "오랜 기간 깔끔하게 착용할 수 있습니다.\n\n"
        "[fit 예시]\n"
        "FREE 사이즈로 77까지 추천드리며\n"
        "여유로운 핏으로 체형 구애 없이 착용 가능합니다.\n"
        "힙을 덮는 기장과 적당한 품으로\n"
        "편안하게 입으실 수 있습니다.\n\n"
        "[occasion 예시]\n"
        "오피스룩, 모임, 데일리 외출 등\n"
        "격식 있는 자리부터 편안한 일상까지\n"
        "다양하게 활용할 수 있는 아이템입니다.\n"
        "시즌에 구애 없이 자주 손이 가는 아이템입니다.\n"
        "</style_guide>\n\n"
        "[금지] 한 줄에 두 문장을 쉼표·접속어로 이어 붙이지 않는다. 한 줄에 30자 초과 금지.\n""[금지] 줄 끝을 쉼표(,)로 절대 끝내지 않는다. 연결되는 줄이라도 쉼표 없이 자연스럽게 끊는다.\n""  예) 금지: 가볍고 내구성이 뛰어나며,  →  허용: 가볍고 내구성이 뛰어나\n""[금지] 마침표와 쉼표가 함께 붙는 문장 금지. 예) 블라우스로,. → 블라우스입니다.\n\n"
        "[입력 데이터]\n"
        f"- 상품명: {data['display_name']}\n"
        f"- 컬러: {data['color']}\n"
        f"- 사이즈: {data['size']}\n"
        f"- 소재: {data['material']}\n"
        f"- 소재설명 참고메모: {data['material_desc_raw']}\n"
        f"- 디테일 특징: {data['detail_tip']}\n"
        f"- 핏/실루엣: {data['fit']}\n"
        f"- 주요 어필 포인트: {data['appeal_points']}\n"
        f"- 기타 특징: {data['etc']}\n"
        f"- 타겟: {data['target']}\n"
        f"- 세탁방법: {data['washing']}\n"
        f"- 실측사이즈: {measurement_str}\n"
        f"- 추가/수정 요청사항: {additional_request or '없음'}\n\n"
        "[JSON 스키마]\n"
        "{\n"
        '  "material_desc_lines": ["문장1.", "문장2.", "문장3."],\n'
        '  "recommend_lines": ["~을 찾으시는 분", "~을 원하시는 분", "~하고 싶으신 분", "~을 중요하게 생각하시는 분"],\n'
        '  "review_lines": ["후기문장1", "후기문장2", "후기문장3", "후기문장4"],\n'
        '  "faqs": [\n'
        '    {"q": "Q. 질문1", "a": "A. 답변1"},\n'
        '    {"q": "Q. 질문2", "a": "A. 답변2"},\n'
        '    {"q": "Q. 질문3", "a": "A. 답변3"},\n'
        '    {"q": "Q. 질문4", "a": "A. 답변4"}\n'
        '  ],\n'
        '  "shopping_lines": ["참고사항1", "참고사항2", "참고사항3"],\n'
        '  "md_sections": {\n'
        '    "choice": ["줄1", "줄2", "줄3", "줄4", "줄5", "줄6"],\n'
        '    "fabric": ["줄1", "줄2", "줄3", "줄4", "줄5", "줄6"],\n'
        '    "fit": ["줄1", "줄2", "줄3", "줄4", "줄5", "줄6"],\n'
        '    "occasion": ["줄1", "줄2", "줄3", "줄4", "줄5", "줄6"],\n'
        '    "purchase_note": [],\n'
        '    "ending": []\n'
        '  },\n'
        '  "size_tips": {\n'
        '    "55": ["체형 착용 특징 첫 문장.", "구체적인 핏·실루엣 설명 둘째 문장."],\n'
        '    "66": ["체형 착용 특징 첫 문장.", "구체적인 핏·실루엣 설명 둘째 문장."],\n'
        '    "66half": ["체형 착용 특징 첫 문장.", "구체적인 핏·실루엣 설명 둘째 문장."],\n'
        '    "77": ["체형 착용 특징 첫 문장.", "구체적인 핏·실루엣 설명 둘째 문장."]\n'
        '  }\n'
        "}\n\n"
        "[각 항목 작성 기준]\n"
        "- material_desc_lines: 아래 규칙을 반드시 지킨다.\n"
        "  * 3문장. 각 문장은 완결된 형태로 마침표로 끝난다.\n"
        "  * 소재의 혼용 구성 또는 소재명 → 표면 질감·광택·두께감 → 착용 시 편안함·관리 편의성 순으로 자연스럽게 서술한다.\n"
        "  * 입력된 소재명과 소재설명 참고메모를 바탕으로 재작성한다. 입력 문구를 그대로 쓰지 않는다.\n"
        "  * 문법 오류 절대 금지. '~가 돋보이는.' 처럼 서술어 없이 끝나는 문장 금지.\n"
        "  * 올바른 예시:\n"
        "    '울, 텐셀, 레이온, 나일론이 조화롭게 혼방되어 부드러운 촉감이 느껴집니다.'\n"
        "    '은은한 광택이 더해져 세련된 무드를 연출하며, 구김이 적어 관리가 편리합니다.'\n"
        "    '가볍고 부담 없는 두께로 여리한 실루엣이 자연스럽게 표현됩니다.'\n"
        "  * 금지 예시: '깨끗한 아이보리 컬러가 돋보이는.' (서술어 없음, 컬러 설명은 소재설명에 불필요)\n\n"
        "- recommend_lines: 반드시 위 추천 블록 규칙 준수. 4줄.\n"
        "- review_lines: 실제 미샵 스텝·모델·MD가 착용 후 말할 법한 솔직하고 구체적인 후기 4줄. 따옴표 없이 문장만.\n"
        "- faqs: 아래 규칙을 반드시 지킨다.\n"
        "  * q는 Q.로 시작, a는 A.로 시작.\n"
        "  * 미샵 4050 여성 고객이 이 상품을 구매 전 실제로 망설이며 물어볼 만한 질문이어야 한다.\n"
        "  * 상품의 소재·핏·사이즈·컬러·디테일·관리법·착용 상황 등 구체적 정보에 기반한 질문과 답변을 작성한다.\n"
        "  * 아래처럼 해당 상품 특성이 반영된 질문을 만든다:\n"
        "    - 체형 관련: '가슴이 있는 77 체형인데 답답하지 않게 입을 수 있을까요?'\n"
        "    - 소재 관련: '큐브라 폴리 소재인데 여름에도 입을 수 있나요?', '비침이 있나요?'\n"
        "    - 컬러 관련: '아이보리라 비침이 심하진 않나요?', '세탁 후 변색되지 않나요?'\n"
        "    - 활용 관련: '오피스룩으로 입기에 너무 캐주얼하지 않나요?'\n"
        "    - 관리 관련: '드라이클리닝만 가능한가요? 손세탁은 안 되나요?'\n"
        "  * 답변은 단답형이 아니라 고객의 걱정을 실제로 해소해주는 충실한 내용으로 작성한다.\n"
        "  * 일반적인 질문 금지 예시: 'Q. 사이즈는 어떻게 선택하나요?' (너무 범용적)\n\n"
        "- shopping_lines: 실용 정보 3줄 (사이즈, 컬러별 주의, 관리방법 등).\n"
        "- md_sections.choice: 이 상품을 선택해야 하는 이유. 자연스럽게 이어지는 4~6줄. 한 줄 15~25자.\n"
        "- md_sections.fabric: 소재와 두께감. 자연스럽게 이어지는 4~6줄. 한 줄 15~25자.\n"
        "- md_sections.fit: 체형, 사이즈 가이드, 실루엣. 자연스럽게 이어지는 4~6줄. 한 줄 15~25자.\n"
        "- md_sections.occasion: 착용 장면과 코디 활용. 자연스럽게 이어지는 4~6줄. 한 줄 15~25자.\n"
        "- size_tips: 각 체형별 2문장. 첫 문장은 착용감·핏 특징, 둘째 문장은 실루엣·커버 효과.\n"
        "  예시) 55: ['여유 있는 핏으로 루즈하게 연출됩니다.', '소매와 어깨가 자연스럽게 떨어져 여리여리한 실루엣이 살아납니다.']\n"
        "  예시) 77: ['상체가 크거나 군살이 있으신 분도 드라마틱한 체형 커버 효과를 느끼실 수 있습니다.', '적당히 여유 있는 핏으로 깔끔하게 연출됩니다.']\n"
    )
    return prompt


# ─────────────────────────────────────────────
# LLM 생성
# ─────────────────────────────────────────────
def generate_structured_copy(data, additional_request, uploaded_images):
    prompt = build_generation_prompt(data, additional_request)
    user_content = [{"type": "text", "text": prompt}]
    for img in (uploaded_images[:5] if uploaded_images else []):
        user_content.append(file_to_content_item(img))

    response = chat_with_retry(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "사용자의 추가/수정 요청사항을 최우선으로 반영한다."},
            {
                "role": "system",
                "content": (
                    "반드시 JSON만 출력한다. 없는 디테일은 추정하지 않는다. "
                    "recommend_lines는 반드시 '~분' 또는 '~분.'으로 끝나는 형태로만 생성한다. "
                    "추천합니다, 권해드립니다, 적합합니다, 분께 같은 표현은 절대 사용하지 않는다. "
                    "md_sections의 각 배열 줄은 절대 쉼표(,)로 끝내지 않는다. "
                    "'블라우스로,' '뛰어나며,' '소재로,' 처럼 쉼표로 끝나는 줄은 완전히 금지한다. "
                    "줄이 다음 줄과 문맥상 이어지더라도 쉼표 없이 자연스럽게 끊거나 완결 문장으로 쓴다."
                ),
            },
            {"role": "user", "content": user_content},
        ],
        temperature=0.55,
        max_retries=2,
    )
    return extract_json(response.choices[0].message.content)


# ─────────────────────────────────────────────
# 폴백 데이터
# ─────────────────────────────────────────────
def fallback_structured(data):
    product = data.get("display_name") or "상품"
    size = clean_line(data.get("size") or "FREE 사이즈로 77까지 추천드립니다.")
    material_desc_raw = data.get("material_desc_raw") or ""
    raw_lines = [x.strip() for x in material_desc_raw.replace("\r", "").split("\n") if x.strip()]
    material_lines = [ensure_period(clean_line(x)) for x in raw_lines[:3]]
    if not material_lines:
        material_lines = [
            "부드러운 터치감으로 피부에 닿는 느낌이 편안한 소재입니다.",
            "표면감이 차분하게 정리되어 데일리로 활용하기 좋습니다.",
            "관리 부담이 크지 않아 손이 자주 가는 아이템입니다.",
        ]

    return {
        "material_desc_lines": material_lines,
        "recommend_lines": [
            "출근룩부터 모임룩까지 단정하게 입을 아이템을 찾으시는 분",
            "상체 라인을 부담 없이 정리해 주는 편안한 핏을 원하시는 분",
            "소재감이 주는 고급스러운 분위기를 중요하게 생각하시는 분",
            "데님, 슬랙스, 스커트와 두루 잘 어울리는 상의를 찾으시는 분",
        ],
        "review_lines": [
            "입었을 때 전체 실루엣이 차분하게 정리돼서 손이 자주 가요.",
            "촉감이 부담스럽지 않아 하루 종일 입어도 편안한 느낌이에요.",
            "격식 있는 자리에도 과하지 않게 잘 어울려 활용도가 높아요.",
            "구김 부담이 크지 않아 바쁜 날에도 깔끔하게 입기 좋았어요.",
        ],
        "faqs": [
            {"q": "Q. 가슴이 있는 77 체형도 답답하지 않게 착용할 수 있을까요?", "a": "A. " + size + " 기준으로 안내드리며, 여유 있게 착용하실 수 있도록 설계되어 있습니다. 실측사이즈를 함께 확인하시면 더 정확합니다."},
            {"q": "Q. 밝은 컬러는 비침이 심한 편인가요?", "a": "A. 밝은 컬러 계열은 약간의 비침이 있을 수 있어 스킨톤 이너와 함께 착용하시면 더 안정감 있게 입으실 수 있습니다."},
            {"q": "Q. 하루 종일 입으면 구김이 많이 남는 편인가요?", "a": "A. 소재 특성상 비교적 구김이 적어 출근룩이나 모임룩으로도 부담 없이 입기 좋은 편입니다."},
            {"q": "Q. 세탁은 어떻게 하면 좋을까요?", "a": "A. " + ensure_period(clean_line(data.get("washing") or "드라이클리닝, 단독 울세탁, 손세탁 권장. 건조기 사용 금지"))},
        ],
        "shopping_lines": [
            size,
            "밝은 컬러 계열은 스킨톤 이너와 함께 착용하시면 더 안정감 있게 연출하실 수 있습니다.",
            "실측사이즈를 함께 확인하시면 원하는 핏으로 선택하시기 더 수월합니다.",
        ],
        "md_sections": {
            "choice": [
                product + "은 과하지 않으면서도 차분하고 세련된 인상을 만들어 줍니다.",
                "기본에 가까운 디자인일수록 소재와 핏의 차이가 크게 드러나는데, 이 아이템은 그 균형감이 특히 좋습니다.",
                "체형을 부드럽게 감싸주는 실루엣이 안정감 있게 느껴집니다.",
                "데일리와 격식을 오가는 장면에서 활용도가 높아 추천드리기 좋습니다.",
            ],
            "fabric": material_lines + ["두께감 또한 과하게 무겁지 않아 계절감에 맞춰 손쉽게 매치하기 좋습니다."],
            "fit": [
                size + " 기준으로 안내드리며, 전체적으로 답답하지 않게 착용하기 좋은 편입니다.",
                "어깨선과 품이 과하게 붙지 않아 상체 라인이 비교적 편안해 보입니다.",
                "실루엣이 과하게 흐트러지지 않아 단정한 분위기를 유지하기 좋습니다.",
                "실측사이즈를 함께 확인하시면 원하는 핏으로 선택하시기 더 수월합니다.",
            ],
            "occasion": [
                "출근룩처럼 단정함이 필요한 날에도 무리 없이 입기 좋습니다.",
                "모임이나 식사 자리처럼 차려입은 느낌이 필요한 순간에도 자연스럽게 어울립니다.",
                "데님, 슬랙스, 스커트와 두루 매치하기 쉬워 코디 폭이 넓습니다.",
                "아우터 안에 받쳐 입거나 단독으로 연출해도 분위기가 흐트러지지 않습니다.",
            ],
            "purchase_note": [],
            "ending": [],
        },
        "size_tips": {
            "55": [
                "전체적으로 여유 있는 느낌으로 떨어져 단독 착용 시에도 부담 없이 활용하기 좋습니다.",
                "소매와 어깨가 자연스럽게 떨어져 여리여리한 실루엣이 살아납니다.",
            ],
            "66": [
                "품과 실루엣이 안정감 있게 정리되어 데일리부터 모임룩까지 자연스럽게 이어집니다.",
                "전체적으로 넉넉한 품으로 상체 군살 커버에 탁월하며 힙을 살짝 덮는 기장감으로 부담 없습니다.",
            ],
            "66half": [
                "상체를 비교적 편안하게 감싸주어 체형 고민을 덜고 입기 좋습니다.",
                "팔과 어깨가 넓은 체형도 불편함 없이 편안하게 착용되며 슬림한 라인이 연출됩니다.",
            ],
            "77": [
                "답답하게 조이지 않고 여유 있게 착용 가능하여 실측 확인 후 선택하시면 만족도가 높습니다.",
                "상체가 크거나 군살이 있으신 분도 드라마틱한 체형 커버 효과를 느끼실 수 있습니다.",
            ],
        },
    }


# ─────────────────────────────────────────────
# 추천 문장 정규화 (핵심 안전장치)
# ─────────────────────────────────────────────
def normalize_recommend_line(x):
    """
    추천 문장이 반드시 '분'으로 끝나도록 보장.
    이중 '분' 방지, 금지 서술형 제거.
    """
    x = clean_line(x)
    # 마침표 제거
    x = x.rstrip(".")
    # 이미 '분'으로 끝나면 그대로
    if x.endswith("분"):
        return x
    # 금지 서술형 어미 제거
    x = re.sub(r"(추천합니다|권해드립니다|적합합니다|잘 어울립니다|좋습니다|만족하실)$", "", x).strip()
    x = x.rstrip(".")
    if x.endswith("분"):
        return x
    return x + " 분"


def normalize_review_line(x):
    """후기 문장에서 따옴표 제거 (렌더 단계에서 붙임)."""
    x = clean_line(x)
    # 다양한 따옴표 제거
    x = x.replace('"', "").replace("\u201c", "").replace("\u201d", "").strip()
    return x


def normalize_faq(faq_raw, fallback):
    out = []
    for item in (faq_raw if isinstance(faq_raw, list) else []):
        if not isinstance(item, dict):
            continue
        q = clean_line(item.get("q", ""))
        a = clean_line(item.get("a", ""))
        if not q or not a:
            continue
        # Q./A. 중복 제거 후 정규화
        q = re.sub(r"^Q\.\s*Q\.", "Q.", q.strip())
        a = re.sub(r"^A\.\s*A\.", "A.", a.strip())
        if not q.startswith("Q."):
            q = "Q. " + q.lstrip("Q.").strip()
        if not a.startswith("A."):
            a = "A. " + a.lstrip("A.").strip()
        out.append({"q": q, "a": a})
        if len(out) >= 4:
            break
    for fb in fallback:
        if len(out) >= 4:
            break
        if all(fb["q"] != x["q"] for x in out):
            out.append(fb)
    return out[:4]


def safe_lines(raw, count, fallback):
    """일반 필드용: LLM 결과가 count보다 부족하면 폴백으로 보충."""
    items = [clean_line(x) for x in (raw if isinstance(raw, list) else []) if clean_line(x)]
    if len(items) < count:
        for fb in fallback:
            c = clean_line(fb)
            if c and c not in items:
                items.append(c)
            if len(items) >= count:
                break
    return items[:count]


def safe_md_lines(raw, fallback):
    """MD원고 전용: LLM이 생성한 줄만 사용. 아예 없을 때만 폴백 전체 사용. 폴백 혼합 없음."""
    def fix_md_line(x):
        x = clean_line(x)
        # 쉼표로 끝나는 줄: 쉼표 제거 (다음 줄과 이어지는 연결줄은 마침표 없이 그대로)
        if x.endswith(","):
            x = x[:-1].rstrip()
        # ",." 중복 패턴 제거
        x = re.sub(r",\.+$", ".", x)
        return x

    items = [fix_md_line(x) for x in (raw if isinstance(raw, list) else []) if clean_line(x)]
    items = [x for x in items if x]  # 빈 줄 제거
    if not items:
        return [clean_line(x) for x in fallback if clean_line(x)]
    return items


def normalize_generated(result, data):
    fb = fallback_structured(data)

    material_desc_lines = [ensure_period(x) for x in safe_lines(result.get("material_desc_lines"), 3, fb["material_desc_lines"])]
    recommend_lines = [normalize_recommend_line(x) for x in safe_lines(result.get("recommend_lines"), 4, fb["recommend_lines"])]
    review_lines = [normalize_review_line(x) for x in safe_lines(result.get("review_lines"), 4, fb["review_lines"])]
    faqs = normalize_faq(result.get("faqs"), fb["faqs"])
    shopping_lines = [ensure_period(x) for x in safe_lines(result.get("shopping_lines"), 3, fb["shopping_lines"])]

    md_raw = result.get("md_sections") if isinstance(result.get("md_sections"), dict) else {}
    md_sections = {
        # safe_md_lines: LLM 생성 줄만 사용, 쉼표 끝 정리 포함. ensure_period 미적용(연결줄 보존)
        "choice":   safe_md_lines(md_raw.get("choice"),   fb["md_sections"]["choice"]),
        "fabric":   safe_md_lines(md_raw.get("fabric"),   fb["md_sections"]["fabric"]),
        "fit":      safe_md_lines(md_raw.get("fit"),      fb["md_sections"]["fit"]),
        "occasion": safe_md_lines(md_raw.get("occasion"), fb["md_sections"]["occasion"]),
        "purchase_note": [],
        "ending": [],
    }

    tips_raw = result.get("size_tips") if isinstance(result.get("size_tips"), dict) else {}

    def normalize_tip(val, fallback_val):
        """size_tips 값이 문자열이든 배열이든 2개 문장 리스트로 정규화."""
        if isinstance(val, list) and val:
            lines = [ensure_period(clean_line(x)) for x in val if clean_line(x)]
            if not lines:
                lines = fallback_val if isinstance(fallback_val, list) else [ensure_period(fallback_val)]
            return lines[:2] if len(lines) >= 2 else lines
        if isinstance(val, str) and val.strip():
            # 문자열이면 마침표 기준으로 분리 시도
            parts = re.split(r"\.\s+", ensure_period(clean_line(val)))
            parts = [p.strip() for p in parts if p.strip()]
            parts = [p if p.endswith(".") else p + "." for p in parts]
            return parts[:2] if len(parts) >= 2 else parts
        # 폴백
        if isinstance(fallback_val, list):
            return fallback_val
        return [ensure_period(fallback_val)]

    size_tips = {
        "55":     normalize_tip(tips_raw.get("55"),     fb["size_tips"]["55"]),
        "66":     normalize_tip(tips_raw.get("66"),     fb["size_tips"]["66"]),
        "66half": normalize_tip(tips_raw.get("66half"), fb["size_tips"]["66half"]),
        "77":     normalize_tip(tips_raw.get("77"),     fb["size_tips"]["77"]),
    }

    return {
        "material_desc_lines": material_desc_lines,
        "recommend_lines": recommend_lines,
        "review_lines": review_lines,
        "faqs": faqs,
        "shopping_lines": shopping_lines,
        "md_sections": md_sections,
        "size_tips": size_tips,
    }


# ─────────────────────────────────────────────
# HTML 렌더링
# ─────────────────────────────────────────────
def render_text_source(structured):
    # 추천 — normalize_recommend_line 한 번 더 (이중 안전장치)
    rec_lines = "".join(
        "▪ " + normalize_recommend_line(x) + "<br>\n"
        for x in structured["recommend_lines"]
    )

    # 후기 — 따옴표로 감싸기 (중복 없이)
    review_lines = "".join(
        '"' + normalize_review_line(x) + '"<br>\n'
        for x in structured["review_lines"]
    )

    # FAQ — Q./A. 중복 방지
    faq_parts = []
    for idx, faq in enumerate(structured["faqs"]):
        q = re.sub(r"^Q\.\s*Q\.", "Q.", faq["q"].strip())
        a = re.sub(r"^A\.\s*A\.", "A.", faq["a"].strip())
        faq_parts.append(q + "<br>\n")
        faq_parts.append(a + "<br>\n")
        if idx < len(structured["faqs"]) - 1:
            faq_parts.append("<br>\n")
    faq_lines = "".join(faq_parts)

    # 쇼핑 참고
    items = structured["shopping_lines"]
    if items:
        shopping_lines = "".join("▪ " + x + "<br>\n" for x in items[:-1])
        shopping_lines += "▪ " + items[-1]
    else:
        shopping_lines = ""

    return (
        '<div style="text-align:center;">\n'
        '<h3 style="margin-bottom:0;">\n'
        "✓ 이런 분께 추천해요!</h3>\n"
        "<br>\n"
        '<p><span style="font-size:14px; line-height:1.8;">\n'
        + rec_lines
        + "</span></p></div>\n"
        "<br><br><br><br>\n\n"
        '<div style="text-align:center;">\n'
        '<h3 style="margin-bottom:0;">\n'
        "✓ 미리 입어 본 착용후기 (모델/스텝/MD리뷰)</h3>\n"
        "<br>\n"
        '<p><span style="font-size:14px; line-height:1.8;">\n'
        + review_lines
        + "</span></p></div>\n"
        "<br><br><br>\n\n"
        '<div style="text-align:center;">\n'
        '<h3 style="margin-bottom:0;">\n'
        "✓ (FAQ) 이 상품, 이게 궁금해요!</h3>\n"
        "<br>\n"
        '<p><span style="font-size:14px; line-height:1.4;">\n'
        + faq_lines
        + "</span></p></div>\n"
        "<br><br><br><br>\n\n"
        '<div style="text-align:center;">\n'
        '<h3 style="margin-bottom:0;">\n'
        "✓쇼핑에 꼭 참고하세요</h3>\n"
        "<br>\n"
        '<p><span style="font-size:14px; line-height:1.8;">\n'
        + shopping_lines
        + "\n</span></p></div>\n"
        "<br><br><br>"
    )


def render_subsc_html(data, structured):
    """MD원고 — 4개 소제목만. purchase_note/ending 절대 포함 안 함."""
    md = structured["md_sections"]

    def join_md(lines):
        return "".join(x + "<br>\n" for x in lines)

    return (
        '<div id="subsc">\n'
        "<h3>" + data["display_name"] + "</h3>\n"
        "<p>\n"
        '<strong style="font-weight:700 !important;">[이 상품을 초이스한 이유입니다.]</strong><br>\n'
        + join_md(md["choice"])
        + "<br>\n"
        '<strong style="font-weight:700 !important;">[원단과 두께 체감에 대하여]</strong><br>\n'
        + join_md(md["fabric"])
        + "<br>\n"
        '<strong style="font-weight:700 !important;">[체형과 핏, 사이즈 선택 가이드]</strong><br>\n'
        + join_md(md["fit"])
        + "<br>\n"
        '<strong style="font-weight:700 !important;">[이렇게 입는 날이 많아집니다]</strong><br>\n'
        + join_md(md["occasion"])
        + "\n</p></div>"
    )


def build_subtap_html(data, material_desc_lines):
    material_line = format_material_line(data["material"])
    washing = (data["washing"] or "").strip() or "드라이클리닝, 단독 울세탁, 손세탁 권장. 건조기 사용 금지"
    size_tip = (data["size"] or "").strip() or "FREE 사이즈로 77까지 추천드립니다."
    measurement_html = format_measurement_lines(data["measurement_lines"])
    desc_html = "<br>\n\t\t\t\t\t\t\t\t\t".join(material_desc_lines) + "<br>"

    return (
        '<div id="Subtap">\n'
        '\t<div id="header2" role="banner">\n'
        '\t\t<nav class="nav" role="navigation">\n'
        '\t\t\t<ul class="nav__list">\n'
        "\t\t\t\t<li>\n"
        '\t\t\t\t\t<input id="group-1" type="checkbox" hidden="">\n'
        '\t\t\t\t\t<label for="group-1" style="border-top-color: rgb(204, 204, 204); border-top-width: 1px; border-top-style: solid;">\n'
        '\t\t\t\t\t\t<p class="fa fa-angle-right"></p>소재 정보</label>\n'
        '\t\t\t\t\t<ul class="group-list">\n'
        "\t\t\t\t\t\t<li>\n"
        '\t\t\t\t\t\t\t<a href="#">\n'
        "\t\t\t\t\t\t\t\t<h3>소재 : " + material_line + "</h3>\n"
        "\t\t\t\t\t\t\t\t<p>\n"
        "\t\t\t\t\t\t\t\t\t" + desc_html + "\n"
        "\t\t\t\t\t\t\t\t</p>\n"
        "\t\t\t\t\t\t\t\t<h3>세탁방법</h3>\n"
        "\t\t\t\t\t\t\t\t<p>" + washing + "</p>\n"
        "\t\t\t\t\t\t\t</a>\n"
        "\t\t\t\t\t\t</li>\n"
        "\t\t\t\t\t</ul>\n"
        "\t\t\t\t</li>\n"
        "\t\t\t\t<li>\n"
        '\t\t\t\t\t<input id="group-2" type="checkbox" hidden="">\n'
        '\t\t\t\t\t<label for="group-2">\n'
        '\t\t\t\t\t\t<p class="fa fa-angle-right"></p>사이즈 정보</label>\n'
        '\t\t\t\t\t<ul class="group-list gray">\n'
        "\t\t\t\t\t\t<li>\n"
        '\t\t\t\t\t\t\t<a href="#">\n'
        "\t\t\t\t\t\t\t\t<h3>사이즈 TIP</h3>\n"
        "\t\t\t\t\t\t\t\t<p>" + size_tip + "</p>\n"
        "\t\t\t\t\t\t\t\t<h3>길이 TIP</h3>\n"
        "\t\t\t\t\t\t\t\t<p>162-167cm에서는 모델핏을 참고해 주시고,\n"
        "<br> 다리 길이나 체형에 따라 다르지만,\n"
        "<br> 160cm이하에서는 모델의 핏보다 조금 길게\n"
        "<br> 연출됩니다.</p>\n"
        "\t\t\t\t\t\t\t</a>\n"
        "\t\t\t\t\t\t</li>\n"
        "\t\t\t\t\t</ul>\n"
        "\t\t\t\t</li>\n"
        "\t\t\t\t<li>\n"
        '\t\t\t\t\t<input id="group-3" type="checkbox" hidden="">\n'
        '\t\t\t\t\t<label for="group-3">\n'
        '\t\t\t\t\t\t<p class="fa fa-angle-right"></p>실측 사이즈</label>\n'
        '\t\t\t\t\t<ul class="group-list">\n'
        "\t\t\t\t\t\t<li>\n"
        '\t\t\t\t\t\t\t<a href="#">\n'
        "\t\t\t\t\t\t\t\t<p>" + measurement_html + "</p>\n"
        "\t\t\t\t\t\t\t</a>\n"
        "\t\t\t\t\t\t</li>\n"
        "\t\t\t\t\t</ul>\n"
        "\t\t\t\t</li>\n"
        "\t\t\t\t<li>\n"
        '\t\t\t\t\t<input id="group-5" type="checkbox" hidden="">\n'
        '\t\t\t\t\t<label for="group-5"><span class="fa fa-angle-right"></span>\n'
        '\t\t\t\t\t\t<a href="#crema-product-fit-1" style="padding: 0px; box-shadow:none; background:#f7f7f7;">실측사이즈 재는방법</a></label>\n'
        "\t\t\t\t</li>\n"
        "\t\t\t</ul>\n"
        "\t\t</nav>\n"
        "\t</div>\n"
        "</div>"
    )


def assemble_final_output(data, structured):
    material_line = format_material_line(data["material"])
    text_source = render_text_source(structured)
    subsc_html = render_subsc_html(data, structured)
    subtap_html = build_subtap_html(data, structured["material_desc_lines"])
    source_block = FIXED_HTML_HEAD + "\n\n" + subsc_html + "\n\n" + subtap_html

    # 소재설명: 단락 형태 (리스트 아님)
    material_desc_paragraph = " ".join(structured["material_desc_lines"])

    lines = []
    lines.append("상품명 : " + data["display_name"])
    lines.append("")
    lines.append("컬러 : " + data["color"])
    lines.append("사이즈 : " + data["size"])
    lines.append("소재 : " + material_line)
    lines.append("소재설명 : " + material_desc_paragraph)
    lines.append("제조국 : " + data["country"])
    lines.append("")
    lines.append("○ 포인트 코멘트")
    lines.append("-")
    # 포인트 코멘트 빈 줄 9개
    for _ in range(9):
        lines.append("")
    lines.append("-")
    lines.append("")
    lines.append("---------------------------------")
    lines.append("텍스트 소스")
    lines.append("---------------------------------")
    lines.append("")
    lines.append(text_source)
    lines.append("")
    lines.append("----------------------------------")
    lines.append("MD원고(상품 설명 소스)")
    lines.append("----------------------------------")
    lines.append(source_block)
    lines.append("")
    lines.append("-----------------")
    lines.append("사이즈 팁")
    lines.append("-----------------")
    lines.append("")
    lines.append("ㅇ55 (90) 160cm 48kg")
    for s in structured["size_tips"]["55"]:
        lines.append(s)
    lines.append("")
    lines.append("ㅇ66 (95) 165cm 54kg")
    for s in structured["size_tips"]["66"]:
        lines.append(s)
    lines.append("")
    lines.append("ㅇ66반 (95) 164cm 58kg")
    for s in structured["size_tips"]["66half"]:
        lines.append(s)
    lines.append("")
    lines.append("ㅇ77 (100) 163cm 61kg")
    for s in structured["size_tips"]["77"]:
        lines.append(s)
    return "\n".join(lines).strip()


# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────
st.markdown("---")
st.subheader("상품 네이밍")
ncol1, ncol2 = st.columns([5, 1], vertical_alignment="bottom")
with ncol1:
    naming_input = st.text_area(
        "상품 주요특징 입력",
        height=120,
        placeholder="예: 여리핏, 부드러운 엠보 텍스처, 상체 군살 커버, 루즈핏 맨투맨",
        key="naming_input_value",
    )
with ncol2:
    if st.button("네이밍 생성", use_container_width=True):
        if naming_input.strip():
            with st.spinner("상품명을 생성 중입니다..."):
                try:
                    response = chat_with_retry(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": NAME_PROMPT},
                            {"role": "user", "content": naming_input},
                        ],
                        temperature=0.5,
                        max_retries=2,
                    )
                    st.session_state.naming_result = response.choices[0].message.content.strip()
                    st.rerun()
                except RateLimitError:
                    st.error("현재 OpenAI 요청이 일시적으로 몰려 네이밍 생성을 완료하지 못했습니다. 잠시 후 다시 시도해 주세요.")
                except Exception as e:
                    st.error(f"네이밍 생성 중 오류가 발생했습니다: {e}")
        else:
            st.warning("상품 주요특징을 입력해 주세요.")

st.text_area("상품 네이밍 제안 20개", value=st.session_state.naming_result, height=250)
st.markdown("---")

h1, h2 = st.columns([2.2, 1.0], vertical_alignment="bottom")
with h1:
    st.subheader("상품정보 입력")
with h2:
    st.button("초기화", use_container_width=True, on_click=reset_all)

nonce = st.session_state.reset_nonce
left, right = st.columns(2)

with left:
    product_name = st.text_input("상품명", value="( color)", key=f"product_name_{nonce}")
    color = st.text_input("컬러", placeholder="예: 1 아이보리 2 연그레이 3 블랙", key=f"color_{nonce}")
    size = st.text_area("사이즈", height=90, value="FREE 사이즈로 77까지 추천드립니다.", key=f"size_{nonce}")
    material = st.text_input("소재", placeholder="예: 울6 텐셀45 레이온37 나일론12", key=f"material_{nonce}")
    material_desc = st.text_area(
        "소재설명",
        height=110,
        placeholder="예: 부드러운 촉감, 은은한 광택, 구김 적음, 여리한 실루엣",
        key=f"material_desc_{nonce}",
    )
    country = st.text_input("제조국", key=f"country_{nonce}")
    top_measure = st.text_area(
        "상의 실측사이즈",
        height=120,
        value="어깨단면 / 가슴둘레 / 암홀둘레 / 소매길이 / 소매둘레 / 총장 / 총장(앞) / 총장(뒤)  단위:cm",
        key=f"top_measure_{nonce}",
    )
    bottom_measure = st.text_area(
        "하의 실측사이즈",
        height=110,
        value="F 허리둘레 / 엉덩이둘레 / 허벅지둘레 / 밑단둘레 / 총장 / 밑위 길이  단위:cm\nL 허리둘레 / 엉덩이둘레 / 허벅지둘레 / 밑단둘레 / 총장 / 밑위 길이  단위:cm",
        key=f"bottom_measure_{nonce}",
    )
    dress_measure = st.text_area(
        "원피스 실측사이즈",
        height=130,
        value="어깨단면 / 가슴둘레 / 허리둘레 / 엉덩이둘레 / 암홀둘레 / 소매길이 / 어깨소매길이 / 총장(앞) / 총장(뒤)  단위:cm",
        key=f"dress_measure_{nonce}",
    )

with right:
    detail_tip = st.text_input("디테일 특징", placeholder="예: 브이넥, 타이 탈부착, 앞 절개라인, 소매 볼륨", key=f"detail_tip_{nonce}")
    fit = st.text_input("핏/실루엣", placeholder="예: 군살커버, 세련된 핏, 여리한 실루엣", key=f"fit_{nonce}")
    appeal_points = st.text_area(
        "주요 어필 포인트",
        height=150,
        placeholder="예: 구김 적음, 체형커버, 오피스룩, 하객룩, 데일리룩",
        key=f"appeal_points_{nonce}",
    )
    etc = st.text_area(
        "기타 특징",
        height=120,
        placeholder="예: 브랜드 퀄리티 고급소재",
        key=f"etc_{nonce}",
    )
    target = st.text_input("타겟", value="4050 여성", key=f"target_{nonce}")
    washing = st.text_input(
        "세탁방법",
        value="드라이클리닝, 단독 울세탁, 손세탁 권장. 건조기 사용 금지",
        key=f"washing_{nonce}",
    )
    additional_request = st.text_area(
        "추가/수정 요청사항 (출력물 확인 후 수정사항 입력)",
        height=120,
        key=f"additional_request_{nonce}",
    )

st.subheader("이미지 업로드")
uploaded_images = st.file_uploader(
    "이미지",
    type=["jpg", "jpeg", "png", "webp"],
    accept_multiple_files=True,
    key=f"uploaded_images_{nonce}",
)

if st.button("생성하기", type="primary", use_container_width=True, key=f"generate_{nonce}"):
    display_name = apply_color_count_to_name(product_name, color)
    measurement_lines = combine_measurements(top_measure, bottom_measure, dress_measure)
    data = {
        "product_name": product_name,
        "display_name": display_name,
        "color": color,
        "size": size,
        "material": material,
        "material_desc_raw": material_desc,
        "country": country,
        "measurement_lines": measurement_lines,
        "detail_tip": detail_tip,
        "fit": fit,
        "appeal_points": appeal_points,
        "etc": etc,
        "target": target,
        "washing": washing,
    }

    with st.spinner("출력물을 생성 중입니다..."):
        try:
            structured_raw = generate_structured_copy(data, additional_request, uploaded_images)
            structured = normalize_generated(structured_raw, data)
            result = assemble_final_output(data, structured)
        except RateLimitError:
            st.error("현재 OpenAI 요청이 일시적으로 몰려 원고 생성을 완료하지 못했습니다. 잠시 후 다시 시도해 주세요.")
            st.stop()
        except Exception as e:
            st.error(f"원고 생성 중 오류가 발생했습니다: {e}")
            st.stop()

    # 세션에 저장 — TXT/HWP 버튼이 다운로드 후에도 사라지지 않음
    st.session_state.generated_result = result
    st.session_state.generated_docx = result_to_docx_bytes(result)
    st.session_state.generated_file_stem = (display_name or "page_builder").replace(" ", "_")

# ─── 결과 표시 (세션 기반) ───
if st.session_state.generated_result:
    st.text_area("결과", st.session_state.generated_result, height=1200)
    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "TXT 다운로드",
            data=st.session_state.generated_result,
            file_name=st.session_state.generated_file_stem + "_output.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with c2:
        st.download_button(
            "HWP 다운로드",
            data=st.session_state.generated_docx,
            file_name=st.session_state.generated_file_stem + "_output.hwp",
            mime="application/x-hwp",
            use_container_width=True,
        )

st.markdown("---")
st.markdown("© made by MISHARP, MIYAWA. All rights reserved.")
